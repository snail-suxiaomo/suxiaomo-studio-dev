"""prompt_library/service.py —— 个人指令库 CRUD

与 novel_prompt_config（系统管线配置）完全独立，这里是用户自己的素材库。
分类字段（category_1 / category_2 / output_type / tags）均为普通文本，
支持「预设 + 用户自建」——前端用 datalist 给建议值，但允许自由输入新值，无需改表。
"""

import json
import os
import io
import re
import shutil
import subprocess
import tarfile
import tempfile
import uuid
import zipfile
from datetime import datetime
from pathlib import Path

from common import db


def _conn():
    return db.get_conn()


# 模版图：真实文件落 data/prompt_images/，库只存相对路径数组（JSON 字符串）
IMAGES_DIR_NAME = 'prompt_images'
ALLOWED_IMG_EXT = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}

# skills 附件：压缩包原文件落 data/prompt_attachments/，md/txt/docx 读取填正文不落库
ATTACHMENTS_DIR_NAME = 'prompt_attachments'
ALLOWED_ATTACH_EXT = {'.zip', '.7z', '.rar', '.tar', '.tgz', '.tar.gz', '.gz'}


def _images_dir():
    d = db.DATA_DIR / IMAGES_DIR_NAME
    d.mkdir(parents=True, exist_ok=True)
    return d


def _attachments_dir(pid=None):
    base = db.DATA_DIR / ATTACHMENTS_DIR_NAME
    if pid is not None:
        base = base / str(pid)
    base.mkdir(parents=True, exist_ok=True)
    return base


# ---- owner / 命名空间 uid ----
# 出厂库 owner_id 固定前缀，保证我们分发的提示词 uid 稳定（v2 重导能跳过、不重复新增）。
# 接收方首次运行生成独立 U-<8位> 前缀存 app_kv，天然不与我们碰撞。
OWNER_ID_DEV = "OEM-7F3A"
OWNER_NAME_DEV = "苏小沫"


def _kv_get(key, default=None):
    conn = _conn()
    try:
        row = conn.execute("SELECT v FROM app_kv WHERE k=?", (key,)).fetchone()
        return row["v"] if row else default
    finally:
        conn.close()


def _kv_set(key, value):
    conn = _conn()
    try:
        conn.execute(
            "INSERT INTO app_kv(k, v) VALUES(?, ?) "
            "ON CONFLICT(k) DO UPDATE SET v=excluded.v", (key, value))
        conn.commit()
    finally:
        conn.close()


def _is_authoring_env():
    p = str(db.DATA_DIR).replace("\\", "/")
    return ("suxiaomo-studio-dev" in p) or ("suxiaomo-studio-workspace" in p)


def get_owner_id():
    oid = _kv_get("owner_id")
    if oid:
        return oid
    oid = OWNER_ID_DEV if _is_authoring_env() else ("U-" + uuid.uuid4().hex[:8])
    _kv_set("owner_id", oid)
    return oid


def get_owner_name():
    on = _kv_get("owner_name")
    if on:
        return on
    on = OWNER_NAME_DEV if _is_authoring_env() else "我"
    _kv_set("owner_name", on)
    return on


def _owner_prefix(uid):
    # uid = owner_id-seq → owner_id 即前缀
    return (uid or "").rsplit("-", 1)[0] if uid else ""


def _next_seq(conn, owner_id):
    prefix = owner_id + "-"
    max_seq = 0
    for r in conn.execute(
            "SELECT uid FROM prompt_library WHERE uid LIKE ?", (prefix + "%",)).fetchall():
        m = re.match(r"^" + re.escape(prefix) + r"(\d+)$", r["uid"] or "")
        if m:
            max_seq = max(max_seq, int(m.group(1)))
    return max_seq + 1


def _parse_images(v):
    if isinstance(v, list):
        return v
    if not v:
        return []
    try:
        a = json.loads(v)
        return a if isinstance(a, list) else []
    except Exception:
        return []


def _dump_images(arr):
    return json.dumps(arr, ensure_ascii=False)


def list_prompts(category='全部', output_type='全部', keyword=None, tag=None, scope='全部', owner_name=None):
    sql = "SELECT * FROM prompt_library WHERE 1=1"
    params = []
    if scope and scope != '全部':
        sql += " AND scope=?"
        params.append(scope)
    if category and category != '全部':
        sql += " AND category=?"
        params.append(category)
    if output_type and output_type != '全部':
        sql += " AND output_type=?"
        params.append(output_type)
    if owner_name:
        sql += " AND owner_name=?"
        params.append(owner_name)
    if tag:
        for t in [x for x in tag.split(',') if x.strip()]:
            sql += " AND tags LIKE ?"
            params.append(f'%{t.strip()}%')
    if keyword:
        like = f'%{keyword}%'
        sql += " AND (title LIKE ? OR content LIKE ? OR note LIKE ? OR tags LIKE ?)"
        params.extend([like, like, like, like])
    sql += " ORDER BY sort_order ASC, updated_at DESC, id DESC"
    conn = _conn()
    try:
        rows = conn.execute(sql, params).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


def get_prompt(pid):
    conn = _conn()
    try:
        row = conn.execute("SELECT * FROM prompt_library WHERE id=?", (pid,)).fetchone()
        return dict(row) if row else None
    finally:
        conn.close()


def create_prompt(data):
    title = (data.get('title') or '').strip()
    if not title:
        raise ValueError('标题不能为空')
    content = data.get('content') or ''
    category = (data.get('category') or '其他').strip() or '其他'
    output_type = (data.get('output_type') or '文本').strip() or '文本'
    note = data.get('note') or ''
    tags = (data.get('tags') or '').strip()
    scope = (data.get('scope') or 'prompt').strip() or 'prompt'
    if data.get('owner_name') is not None:
        owner_name = str(data.get('owner_name')).strip()
        if not owner_name:
            raise ValueError('作者不能为空')
    else:
        owner_name = get_owner_name()
    owner_id = get_owner_id()
    conn = _conn()
    try:
        # 标题同 owner_name 内唯一（注册式）：同一作者下不允许两个同名提示词
        if conn.execute(
                "SELECT 1 FROM prompt_library WHERE owner_name=? AND title=?",
                (owner_name, title)).fetchone():
            raise ValueError(f'同一作者下已存在标题《{title}》，请换一个')
        max_order = conn.execute(
            "SELECT COALESCE(MAX(sort_order), 0) FROM prompt_library"
        ).fetchone()[0]
        new_uid = f"{owner_id}-{_next_seq(conn, owner_id):04d}"
        cur = conn.execute(
            "INSERT INTO prompt_library(title, content, category, category_1, category_2, output_type, note, tags, source_file, images, scope, sort_order, uid, owner_name) "
            "VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (title, content, category, category, '通用', output_type, note, tags, None, '[]', scope, max_order + 10, new_uid, owner_name),
        )
        conn.commit()
        return get_prompt(cur.lastrowid)
    finally:
        conn.close()


def batch_create_prompts(items):
    """批量新建提示词（供「导入外部提示词」使用）。
    items: list of dict，字段 title / content / category / output_type / style / tool / tags / note / owner_name。
    每条 owner_name 缺省取当前用户；风格/工具合入 tags；同作者内重名跳过（不覆盖）。
    返回统计 + 本次新建的 id 列表（供调用方逐条补图）。"""
    stats = {'created': 0, 'skipped': 0, 'errors': [], 'created_ids': []}
    conn = _conn()
    try:
        max_order = conn.execute(
            "SELECT COALESCE(MAX(sort_order), 0) FROM prompt_library"
        ).fetchone()[0]
        for it in items or []:
            title = (it.get('title') or '').strip()
            if not title:
                stats['errors'].append('存在空标题，已跳过')
                continue
            content = it.get('content') or ''
            note = it.get('note') or ''
            owner_name = (it.get('owner_name') or get_owner_name()).strip() or get_owner_name()
            category = (it.get('category') or '其他').strip() or '其他'
            output_type = (it.get('output_type') or '文本').strip() or '文本'
            style = (it.get('style') or '通用风格').strip() or '通用风格'
            tool = (it.get('tool') or '通用工具').strip() or '通用工具'
            tags = (it.get('tags') or f'{style},{tool}').strip()
            # 同作者内重名跳过（与单条新建一致，避免误覆盖）
            if conn.execute(
                    "SELECT 1 FROM prompt_library WHERE owner_name=? AND title=?",
                    (owner_name, title)).fetchone():
                stats['skipped'] += 1
                continue
            owner_id = get_owner_id()
            max_order += 10
            new_uid = f"{owner_id}-{_next_seq(conn, owner_id):04d}"
            cur = conn.execute(
                "INSERT INTO prompt_library(title, content, category, category_1, category_2, output_type, note, tags, source_file, images, scope, sort_order, uid, owner_name) "
                "VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (title, content, category, category, '通用', output_type, note, tags, None, '[]', 'prompt', max_order, new_uid, owner_name),
            )
            stats['created'] += 1
            stats['created_ids'].append(cur.lastrowid)
        conn.commit()
    finally:
        conn.close()
    return stats


def update_prompt(pid, data):
    existing = get_prompt(pid)
    if not existing:
        raise ValueError('记录不存在')
    title = data.get('title')
    if title is not None:
        title = title.strip()
        if not title:
            raise ValueError('标题不能为空')
    else:
        title = existing['title']
    content = data['content'] if data.get('content') is not None else existing['content']
    category = (data.get('category') or existing.get('category') or existing['category_1']).strip() or '其他'
    output_type = (data.get('output_type') or existing['output_type']).strip() or '文本'
    note = data['note'] if data.get('note') is not None else existing['note']
    tags = (data.get('tags') or '').strip()
    scope = data.get('scope') or existing.get('scope') or 'prompt'
    owner_name = (data.get('owner_name') or existing.get('owner_name') or get_owner_name()).strip()
    if not owner_name:
        raise ValueError('作者不能为空')
    conn = _conn()
    try:
        # 改名时校验同作者内唯一（排除自身）；uid 不改（作者归属固定）
        if title != existing['title']:
            if conn.execute(
                    "SELECT 1 FROM prompt_library WHERE owner_name=? AND title=? AND id!=?",
                    (owner_name, title, pid)).fetchone():
                raise ValueError(f'同一作者下已存在标题《{title}》，请换一个')
        conn.execute(
            "UPDATE prompt_library SET title=?, content=?, category=?, category_1=?, category_2=?, output_type=?, note=?, tags=?, scope=?, owner_name=?, "
            "updated_at=(datetime('now','localtime')) WHERE id=?",
            (title, content, category, category, '通用', output_type, note, tags, scope, owner_name, pid),
        )
        conn.commit()
        return get_prompt(pid)
    finally:
        conn.close()


def delete_prompt(pid):
    conn = _conn()
    try:
        # 级联清理 skills 附件（磁盘文件 + DB 行）
        rows = conn.execute(
            "SELECT filepath FROM prompt_attachments WHERE prompt_id=?", (pid,)).fetchall()
        for r in rows:
            p = (db.DATA_DIR / r['filepath']).resolve()
            if p.is_file():
                try:
                    p.unlink()
                except OSError:
                    pass
        conn.execute("DELETE FROM prompt_attachments WHERE prompt_id=?", (pid,))
        conn.execute("DELETE FROM prompt_library WHERE id=?", (pid,))
        conn.commit()
    finally:
        conn.close()


def set_first_image(pid, img):
    """把指定图片移到 images 数组首位（即设为首图），持久化保存。"""
    rec = get_prompt(pid)
    if not rec:
        raise ValueError('记录不存在')
    arr = _parse_images(rec.get('images')) or []
    if img not in arr:
        raise ValueError('图片不存在')
    arr.remove(img)
    arr.insert(0, img)
    conn = _conn()
    try:
        conn.execute(
            "UPDATE prompt_library SET images=?, updated_at=(datetime('now','localtime')) WHERE id=?",
            (_dump_images(arr), pid),
        )
        conn.commit()
        return get_prompt(pid)
    finally:
        conn.close()


def reorder_prompts(order_ids):
    """按 id 顺序写 sort_order（0..n-1），忽略不存在 id、缺失补末尾。"""
    conn = _conn()
    try:
        valid = {r[0] for r in conn.execute(
            "SELECT id FROM prompt_library").fetchall()}
        filtered = [int(x) for x in order_ids if int(x) in valid]
        # 把没出现的 id 按当前 sort_order 追加到末尾，避免遗漏
        existing_sorted = [r[0] for r in conn.execute(
            "SELECT id FROM prompt_library ORDER BY sort_order ASC, id ASC").fetchall()]
        tail = [pid for pid in existing_sorted if pid not in filtered]
        final = filtered + tail
        for idx, pid in enumerate(final):
            conn.execute(
                "UPDATE prompt_library SET sort_order=? WHERE id=?",
                (idx, pid))
        conn.commit()
    finally:
        conn.close()


def add_images(pid, files):
    """files: list of {filename, content(bytes)}；保存到 data/prompt_images/，追加到 images 数组。"""
    rec = get_prompt(pid)
    if not rec:
        raise ValueError('记录不存在')
    arr = _parse_images(rec.get('images'))
    d = _images_dir()
    for f in files:
        ext = Path(f.get('filename') or '').suffix.lower()
        if ext not in ALLOWED_IMG_EXT:
            raise ValueError(f'不支持的图片格式：{ext or "(无扩展名)"}')
        name = f"{pid}_{uuid.uuid4().hex[:8]}" + ext
        (d / name).write_bytes(f.get('content') or b'')
        arr.append(f"{IMAGES_DIR_NAME}/{name}")
    conn = _conn()
    try:
        conn.execute(
            "UPDATE prompt_library SET images=?, updated_at=(datetime('now','localtime')) WHERE id=?",
            (_dump_images(arr), pid),
        )
        conn.commit()
        return get_prompt(pid)
    finally:
        conn.close()


def delete_image(pid, filename):
    """删除单张模版图（从数组移除 + 删磁盘文件）。filename 为纯文件名（不含目录）。"""
    rec = get_prompt(pid)
    if not rec:
        raise ValueError('记录不存在')
    arr = _parse_images(rec.get('images'))
    rel = f"{IMAGES_DIR_NAME}/{filename}"
    if rel not in arr:
        raise ValueError('图片不存在')
    arr.remove(rel)
    fp = _images_dir() / filename
    if fp.exists():
        try:
            fp.unlink()
        except OSError:
            pass
    conn = _conn()
    try:
        conn.execute(
            "UPDATE prompt_library SET images=?, updated_at=(datetime('now','localtime')) WHERE id=?",
            (_dump_images(arr), pid),
        )
        conn.commit()
        return get_prompt(pid)
    finally:
        conn.close()


def get_image_file(rel_path):
    """返回图片绝对路径，限制在 data/prompt_images 内（防目录穿越）。"""
    p = (db.DATA_DIR / rel_path).resolve()
    base = _images_dir().resolve()
    if p != base and base not in p.parents:
        raise ValueError('非法路径')
    if not p.is_file():
        raise ValueError('图片不存在')
    return p


# ---------------------------------------------------------------------------
# skills 附件：压缩包原文件（md/txt/docx 走读取填正文，不落附件）
# ---------------------------------------------------------------------------

def list_attachments(pid):
    conn = _conn()
    try:
        rows = conn.execute(
            "SELECT id, prompt_id, filename, filetype, filesize, created_at "
            "FROM prompt_attachments WHERE prompt_id=? ORDER BY id",
            (pid,)).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


def add_attachments(pid, files):
    """files: list of {filename, content(bytes)}；仅压缩类扩展名，保存到 data/prompt_attachments/{pid}/。
    前提：只有 category='skills' 的提示词才有附件，其他分类一律拒绝（与前端 v-if 一致）。"""
    rec = get_prompt(pid)
    if not rec:
        raise ValueError('记录不存在')
    if (rec.get('category') or '') != 'skills':
        raise ValueError('仅 skills 分类支持附件')
    saved = []
    d = _attachments_dir(pid)
    conn = _conn()
    try:
        for f in files:
            name = f.get('filename') or ''
            ext = Path(name).suffix.lower()
            if ext not in ALLOWED_ATTACH_EXT:
                raise ValueError(f'仅支持压缩包（zip/7z/rar/tar 等），不支持：{ext or "(无扩展名)"}')
            # 防目录穿越：只用纯文件名
            pure = Path(name).name
            disk_name = f"{uuid.uuid4().hex[:8]}_{pure}"
            content = f.get('content') or b''
            (d / disk_name).write_bytes(content)
            rel = f"{ATTACHMENTS_DIR_NAME}/{pid}/{disk_name}"
            cur = conn.execute(
                "INSERT INTO prompt_attachments(prompt_id, filename, filetype, filesize, filepath) "
                "VALUES(?,?,?,?,?)",
                (pid, pure, ext.lstrip('.'), len(content), rel))
            saved.append({'id': cur.lastrowid, 'filename': pure, 'filetype': ext.lstrip('.'), 'filesize': len(content)})
        conn.commit()
        return saved
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def get_attachment_file(pid, aid):
    """返回附件绝对路径，限制在 data/prompt_attachments/{pid}/ 内。"""
    conn = _conn()
    try:
        row = conn.execute(
            "SELECT * FROM prompt_attachments WHERE id=? AND prompt_id=?", (aid, pid)).fetchone()
    finally:
        conn.close()
    if not row:
        raise ValueError('附件不存在')
    p = (db.DATA_DIR / row['filepath']).resolve()
    base = _attachments_dir(pid).resolve()
    if p != base and base not in p.parents:
        raise ValueError('非法路径')
    if not p.is_file():
        raise ValueError('附件文件不存在')
    return p, dict(row)


def delete_attachment(pid, aid):
    conn = _conn()
    try:
        row = conn.execute(
            "SELECT * FROM prompt_attachments WHERE id=? AND prompt_id=?", (aid, pid)).fetchone()
        if not row:
            raise ValueError('附件不存在')
        p = (db.DATA_DIR / row['filepath']).resolve()
        if p.is_file():
            try:
                p.unlink()
            except OSError:
                pass
        conn.execute("DELETE FROM prompt_attachments WHERE id=?", (aid,))
        conn.commit()
        return {'ok': True}
    finally:
        conn.close()


def extract_docx_text(file_bytes: bytes):
    """docx → 纯文本（python-docx）。返回段落文本（\n 连接）。"""
    from docx import Document  # python-docx
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if (c.text or '').strip()]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


def meta(scope='全部'):
    """返回已有的分类/形态/标签，供前端 chips 与下拉使用。可按 scope 过滤。"""
    conn = _conn()
    try:
        scope_cond = ""
        sp = []
        if scope and scope != '全部':
            scope_cond = " AND scope=?"
            sp = [scope]
        cats = [r[0] for r in conn.execute(
            "SELECT DISTINCT category FROM prompt_library WHERE category IS NOT NULL AND category<>''"
            + scope_cond + " ORDER BY category", sp).fetchall()]
        ot = [r[0] for r in conn.execute(
            "SELECT DISTINCT output_type FROM prompt_library WHERE output_type IS NOT NULL AND output_type<>''"
            + scope_cond + " ORDER BY output_type", sp).fetchall()]
        tags = set()
        for r in conn.execute(
                "SELECT DISTINCT tags FROM prompt_library WHERE tags IS NOT NULL AND tags<>''"
                + scope_cond, sp).fetchall():
            for t in (r[0] or '').split(','):
                t = t.strip()
                if t:
                    tags.add(t)
        owners = [r[0] for r in conn.execute(
            "SELECT DISTINCT owner_name FROM prompt_library WHERE owner_name IS NOT NULL AND owner_name<>''"
            + scope_cond + " ORDER BY owner_name", sp).fetchall()]
        return {
            'categories': cats,
            'output_types': ot,
            'all_tags': sorted(tags),
            'owners': owners,
        }
    finally:
        conn.close()


# ============ skills 分类：上传压缩包并固定落盘到 data/skills/ ============
SKILLS_DIR_NAME = 'skills'
ALLOWED_ARCHIVE_EXT = {'.zip', '.7z', '.rar', '.tar', '.tgz', '.tar.gz',
                       '.tar.bz2', '.tar.xz', '.tbz2', '.txz'}
MAX_ARCHIVE_BYTES = 200 * 1024 * 1024          # 压缩包上限 200MB
MAX_FILES = 3000                               # 解压后文件数上限
MAX_EXTRACTED_BYTES = 2 * 1024 * 1024 * 1024   # 解压后总大小上限 2GB


def _skills_dir():
    d = db.DATA_DIR / SKILLS_DIR_NAME
    d.mkdir(parents=True, exist_ok=True)
    return d


def _slug(s):
    """把 skill 名称转成安全的目录名（保留中文/字母/数字，其余变下划线）。"""
    s = (s or '').strip()
    s = re.sub(r'[^\w一-龥\-]+', '_', s)
    s = re.sub(r'_+', '_', s).strip('_')
    return s or 'skill'


def _find_7z():
    """查找可用的 7-Zip 可执行文件：打包态预留 tools/7z/，否则 PATH。"""
    cand = [
        Path(__file__).resolve().parent.parent / 'tools' / '7z' / '7z.exe',
        Path(__file__).resolve().parent.parent / 'tools' / '7z' / '7za.exe',
    ]
    for c in cand:
        if c.exists():
            return str(c)
    for name in ('7z', '7za', '7zr'):
        p = shutil.which(name)
        if p:
            return p
    return None


def _safe_name(base, name):
    """校验解压条目路径，禁止目录穿越（返回用于写入的目标路径）。"""
    target = (base / name).resolve()
    base_r = base.resolve()
    if target != base_r and base_r not in target.parents:
        raise ValueError(f'压缩包含非法路径: {name}')
    return target


def _enforce_limits(dest: Path):
    """解压后统一校验文件数与总大小，超限则清空并抛错。"""
    total = 0
    count = 0
    for root, _dirs, files in os.walk(dest):
        for f in files:
            count += 1
            if count > MAX_FILES:
                shutil.rmtree(dest, ignore_errors=True)
                raise ValueError(f'解压后文件数超过上限 {MAX_FILES}')
            try:
                total += (Path(root) / f).stat().st_size
            except OSError:
                pass
            if total > MAX_EXTRACTED_BYTES:
                shutil.rmtree(dest, ignore_errors=True)
                raise ValueError('解压后总大小超过上限 2GB')
    return count


def _extract_zip(archive: Path, dest: Path):
    with zipfile.ZipFile(archive) as z:
        for n in z.namelist():
            if n.endswith('/') or n.endswith('\\'):
                continue
            _safe_name(dest, n)
        z.extractall(dest)
    _enforce_limits(dest)


def _extract_tar(archive: Path, dest: Path):
    with tarfile.open(archive, 'r:*') as t:
        for m in t.getmembers():
            if m.isfile():
                _safe_name(dest, m.name)
        t.extractall(dest)
    _enforce_limits(dest)


def _extract_7z_py(archive: Path, dest: Path):
    import py7zr
    with py7zr.SevenZipFile(archive, 'r') as z:
        for n in z.getnames():
            if n.endswith('/') or n.endswith('\\'):
                continue
            _safe_name(dest, n)
        z.extractall(path=str(dest))
    _enforce_limits(dest)


def _extract_via_7z(archive: Path, dest: Path):
    exe = _find_7z()
    if not exe:
        raise ValueError(
            '未找到 7-Zip 可用程序。7z/rar 压缩包请先安装 7-Zip（或把 7z.exe 放到 '
            'backend/tools/7z/），或改用 zip / tar 压缩包。')
    dest.mkdir(parents=True, exist_ok=True)
    r = subprocess.run([exe, 'x', '-y', f'-o{dest}', str(archive)],
                       capture_output=True, text=True)
    if r.returncode != 0:
        raise ValueError('解压失败: ' + (r.stderr or r.stdout or '').strip()[:300])
    _enforce_limits(dest)


def extract_archive(archive: Path, dest: Path, ext: str):
    """按扩展名分发到对应解压器。rar 只能走 7-Zip 二进制。"""
    ext = ext.lower()
    if ext == '.zip':
        _extract_zip(archive, dest)
    elif ext in ('.tar', '.tgz', '.tar.gz', '.tar.bz2', '.tar.xz', '.tbz2', '.txz'):
        _extract_tar(archive, dest)
    elif ext == '.7z':
        try:
            _extract_7z_py(archive, dest)
        except ImportError:
            _extract_via_7z(archive, dest)
    elif ext == '.rar':
        _extract_via_7z(archive, dest)
    else:
        raise ValueError(f'不支持的压缩格式：{ext}')


def _archive_ext(filename):
    """取压缩包扩展名：正确识别 .tar.gz / .tar.bz2 / .tar.xz 等双段扩展名。"""
    name = (filename or '').lower()
    for long_ext in ('.tar.gz', '.tar.bz2', '.tar.xz', '.tgz', '.tbz2', '.txz'):
        if name.endswith(long_ext):
            return long_ext
    return Path(filename).suffix.lower()


def upload_skill(title, archive_bytes, filename, note=''):
    if not title:
        raise ValueError('skill 名称不能为空')
    ext = _archive_ext(filename)
    if ext not in ALLOWED_ARCHIVE_EXT:
        raise ValueError(f'仅支持压缩包：zip / 7z / rar / tar 系，当前为 {ext or "(无扩展名)"}')
    if len(archive_bytes) > MAX_ARCHIVE_BYTES:
        raise ValueError('压缩包超过 200MB 上限')

    skills = _skills_dir()
    # 唯一文件夹名：安全名 + 短随机码
    base_slug = _slug(title)
    folder = f'{base_slug}_{uuid.uuid4().hex[:8]}'
    dest = skills / folder
    while dest.exists():
        folder = f'{base_slug}_{uuid.uuid4().hex[:8]}'
        dest = skills / folder
    dest.mkdir(parents=True, exist_ok=True)

    tmp = None
    try:
        # 落临时文件再用标准解压器处理
        fd, tmp_path = tempfile.mkstemp(suffix=ext, prefix='skill_up_', dir=str(skills))
        os.close(fd)
        tmp = Path(tmp_path)
        tmp.write_bytes(archive_bytes)
        extract_archive(tmp, dest, ext)
    except Exception:
        shutil.rmtree(dest, ignore_errors=True)
        raise
    finally:
        if tmp and tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass

    # 建库记录：source_file 存相对路径，便于将来「打开目录」
    rec = _create_skill_record(title=title, folder_rel=f'{SKILLS_DIR_NAME}/{folder}/', note=note)
    return rec


def _create_skill_record(title, folder_rel, note=''):
    conn = _conn()
    try:
        max_order = conn.execute(
            "SELECT COALESCE(MAX(sort_order), 0) FROM prompt_library").fetchone()[0]
        cur = conn.execute(
            "INSERT INTO prompt_library(title, content, category, category_1, category_2, "
            "output_type, note, tags, source_file, images, scope, sort_order) "
            "VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
            (title, note or '', 'skills', 'skills', '通用', '其他', note or '',
             'skill', folder_rel, '[]', 'instruction', max_order + 10),
        )
        conn.commit()
        return get_prompt(cur.lastrowid)
    finally:
        conn.close()


# ============ 导出 / 导入 / 批量删除（自包含快照，保证删后重导一致） ============
def export_records(ids=None, scope=None):
    """导出为自包含 ZIP：每个提示词一个文件夹（以标题命名），内含 prompt.txt 正文 + meta.json（除正文外所有字段，含 uid/owner_name）+ images/ 子目录。
    返回 (zip_bytes, stats)。不传 ids 导出全部；传 scope 则只导出该归属。"""
    conn = _conn()
    try:
        if ids:
            ids = [int(x) for x in ids]
            placeholders = ','.join('?' * len(ids))
            rows = conn.execute(
                f"SELECT * FROM prompt_library WHERE id IN ({placeholders}) ORDER BY sort_order ASC, id ASC",
                ids).fetchall()
        elif scope and scope != '全部':
            rows = conn.execute(
                "SELECT * FROM prompt_library WHERE scope=? ORDER BY sort_order ASC, id ASC",
                (scope,)).fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM prompt_library ORDER BY sort_order ASC, id ASC").fetchall()
        rows = [dict(r) for r in rows]
        # 附件一次性批量取出，按 prompt_id 分组（避免循环内反复新建/关闭连接）
        atts_by_pid = {}
        if rows:
            pids = [r['id'] for r in rows]
            ph = ','.join('?' * len(pids))
            for a in conn.execute(
                f"SELECT * FROM prompt_attachments WHERE prompt_id IN ({ph})", pids).fetchall():
                atts_by_pid.setdefault(a['prompt_id'], []).append(dict(a))
    finally:
        conn.close()

    img_dir = _images_dir()
    skills_dir = _skills_dir()
    missing_images = []
    used_names = {}

    def safe_folder(title):
        raw = (title or '').strip() or '未命名'
        for ch in '\\/:*?"<>|':
            raw = raw.replace(ch, '_')
        raw = raw.strip().rstrip('.') or '未命名'
        if raw not in used_names:
            used_names[raw] = 0
            return raw
        used_names[raw] += 1
        return f'{raw}（{used_names[raw]}）'

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as z:
        for rec in rows:
            r = dict(rec)
            folder = safe_folder(r.get('title'))
            # 正文
            z.writestr(f'{folder}/prompt.txt', r.get('content') or '', compress_type=zipfile.ZIP_DEFLATED)
            # meta：除正文外所有字段（含 uid / owner_name，并推导 owner_id 由 uid 前缀）
            meta = {k: r.get(k) for k in (
                'uid', 'owner_name', 'title', 'category', 'category_1', 'category_2',
                'output_type', 'note', 'tags', 'source_file', 'scope', 'sort_order',
                'created_at', 'updated_at')}
            if r.get('uid'):
                meta['owner_id'] = _owner_prefix(r['uid'])
            # 图片：按数组顺序写进 images/，首图在前；图片本身已压缩，用 STORED 避免浪费 CPU
            imgs = _parse_images(r.get('images'))
            img_names = []
            for idx, rel in enumerate(imgs, 1):
                fname = Path(rel).name
                src = img_dir / fname
                if src.is_file():
                    arc = f'{folder}/images/{idx:02d}_{fname}'
                    z.write(str(src), arc, compress_type=zipfile.ZIP_STORED)
                    img_names.append(fname)
                else:
                    missing_images.append({'title': r.get('title'), 'file': rel})
            meta['_image_files'] = img_names
            # skills 文件夹（边角情况）：整个文件夹打进该提示词下的 skills/
            skill_folder = None
            if r.get('category') == 'skills' and r.get('source_file'):
                folder_rel = str(r['source_file']).lstrip('/')
                src_folder = (db.DATA_DIR / folder_rel).resolve()
                base = skills_dir.resolve()
                if src_folder.is_dir() and (src_folder == base or base in src_folder.parents):
                    skill_folder = folder_rel
                    for root, _dirs, files in os.walk(src_folder):
                        for f in files:
                            fp = Path(root) / f
                            arc = f'{folder}/skills/{fp.relative_to(src_folder).as_posix()}'
                            z.write(str(fp), arc, compress_type=zipfile.ZIP_STORED)
            meta['_skill_folder'] = skill_folder
            # 附件（skills 压缩包）：打进 attachments/，meta 记文件名（已批量预取，直接按 pid 取）
            atts = atts_by_pid.get(r.get('id'), [])
            att_names = []
            for a in atts:
                src = (db.DATA_DIR / a['filepath']).resolve()
                if src.is_file():
                    arc = f"{folder}/attachments/{a['filename']}"
                    z.write(str(src), arc, compress_type=zipfile.ZIP_STORED)
                    att_names.append(a['filename'])
            meta['_attachment_files'] = att_names
            z.writestr(f'{folder}/meta.json', json.dumps(meta, ensure_ascii=False, indent=2), compress_type=zipfile.ZIP_DEFLATED)
        manifest = {
            'version': 2,
            'format': 'folder-per-prompt',
            'exported_at': datetime.now().isoformat(timespec='seconds'),
            'owner_id': get_owner_id(),
            'owner_name': get_owner_name(),
            'count': len(rows),
            'missing_images': missing_images,
        }
        z.writestr('manifest.json', json.dumps(manifest, ensure_ascii=False, indent=2))
    return buf.getvalue(), {'count': len(rows), 'missing_images': missing_images}


def import_records(file_bytes, filename):
    """导入新格式 ZIP（每提示词一文件夹：prompt.txt + meta.json + images/）。
    匹配规则（只新增、不更新、不删除）：优先 uid 精确命中；否则 同作者(owner_name) + 同 title 命中；命中即跳过。
    未命中则新增（保留包内 uid/owner_name，sort_order 追尾到末尾）。返回统计 imported/skipped/failed/errors。"""
    name = (filename or '').lower()
    if not name.endswith('.zip'):
        raise ValueError('仅支持本工具导出的新格式 .zip 包；旧版 JSON / Zip 请重新导出后再导入。')
    folders = {}
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            for n in z.namelist():
                if n.endswith('/'):
                    continue
                parts = n.split('/')
                if len(parts) < 2:
                    continue
                folder, inner = parts[0], '/'.join(parts[1:])
                e = folders.setdefault(folder, {})
                if inner == 'meta.json':
                    e['meta'] = json.loads(z.read(n).decode('utf-8'))
                elif inner == 'prompt.txt':
                    e['content'] = z.read(n).decode('utf-8')
                elif inner.startswith('images/'):
                    e.setdefault('images', {})[parts[-1]] = z.read(n)
                elif inner.startswith('skills/'):
                    e.setdefault('skills', {})[inner[len('skills/'):]] = z.read(n)
    except zipfile.BadZipFile:
        raise ValueError('ZIP 文件已损坏或不是有效的压缩包')
    prompts = [(f, e) for f, e in folders.items() if 'meta' in e]
    if not prompts:
        raise ValueError('压缩包内未找到提示词数据（请用本工具导出的新格式 zip）')

    img_dir = _images_dir()
    skills_dir = _skills_dir()
    stats = {'imported': 0, 'skipped': 0, 'failed': 0, 'errors': []}
    conn = _conn()
    try:
        existing_uids = {r[0] for r in conn.execute(
            "SELECT uid FROM prompt_library WHERE uid IS NOT NULL AND uid <> ''").fetchall()}

        def dup_owner_title(owner_name, title):
            if not owner_name:
                return False
            return conn.execute(
                "SELECT 1 FROM prompt_library WHERE owner_name=? AND title=?",
                (owner_name, title)).fetchone() is not None

        max_order = conn.execute("SELECT COALESCE(MAX(sort_order), 0) FROM prompt_library").fetchone()[0]
        for folder, e in prompts:
            meta = e.get('meta', {})
            uid = (meta.get('uid') or '').strip()
            owner_id = (meta.get('owner_id') or _owner_prefix(uid) or '').strip()
            owner_name = (meta.get('owner_name') or get_owner_name()).strip() or get_owner_name()
            title = (meta.get('title') or '').strip()
            if not title:
                stats['failed'] += 1
                stats['errors'].append(f'「{folder}」缺少标题，已跳过')
                continue
            # 只新增、存在则跳过（不更新，避免误改/重复）
            if uid and uid in existing_uids:
                stats['skipped'] += 1
                continue
            if dup_owner_title(owner_name, title):
                stats['skipped'] += 1
                continue
            category = (meta.get('category') or '其他').strip() or '其他'
            category_1 = meta.get('category_1') or category
            category_2 = meta.get('category_2') or '通用'
            output_type = (meta.get('output_type') or '文本').strip() or '文本'
            note = meta.get('note') or ''
            tags = meta.get('tags') or ''
            scope = meta.get('scope') or 'prompt'
            content = e.get('content', '') or ''
            # 还原图片（按文件名排序，首图在前）
            images_list = []
            for fn in sorted(e.get('images', {}).keys()):
                data = e['images'][fn]
                ext = Path(fn).suffix
                newfn = f'{uuid.uuid4().hex[:10]}{ext}'
                (img_dir / newfn).write_bytes(data)
                images_list.append(f'prompt_images/{newfn}')
            images_json = json.dumps(images_list, ensure_ascii=False)
            # 还原 skills 文件夹（边角情况）
            source_file = None
            if e.get('skills') and category == 'skills':
                folder_rel = f'skills/{owner_id or "u"}_{(uid or "x")[-6:]}'
                dest = (db.DATA_DIR / folder_rel).resolve()
                base = skills_dir.resolve()
                if dest != base and base not in dest.parents and not dest.exists():
                    dest.mkdir(parents=True, exist_ok=True)
                    for rel, data in e['skills'].items():
                        fp = (db.DATA_DIR / folder_rel / rel).resolve()
                        if fp == base or base in fp.parents:
                            continue
                        fp.parent.mkdir(parents=True, exist_ok=True)
                        fp.write_bytes(data)
                    source_file = folder_rel
            max_order += 10
            conn.execute(
                "INSERT INTO prompt_library(uid, owner_name, title, content, category, category_1, category_2, "
                "output_type, note, tags, source_file, images, scope, sort_order) "
                "VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (uid, owner_name, title, content, category, category_1, category_2, output_type,
                 note, tags, source_file, images_json, scope, max_order))
            stats['imported'] += 1
            if uid:
                existing_uids.add(uid)
        conn.commit()
    finally:
        conn.close()
    return stats


def batch_delete(ids):
    """批量删除：同时移除磁盘上的图片与 skills 文件夹，保持库与磁盘一致。"""
    ids = [int(x) for x in ids]
    if not ids:
        return {'deleted': 0}
    placeholders = ','.join('?' * len(ids))
    img_dir = _images_dir()
    skills_dir = _skills_dir()
    conn = _conn()
    try:
        rows = conn.execute(
            f"SELECT * FROM prompt_library WHERE id IN ({placeholders})", ids).fetchall()
        for r in rows:
            r = dict(r)
            for rel in _parse_images(r.get('images')):
                fp = img_dir / Path(rel).name
                if fp.exists():
                    try:
                        fp.unlink()
                    except OSError:
                        pass
            if r.get('category') == 'skills' and r.get('source_file'):
                folder = (db.DATA_DIR / str(r['source_file']).lstrip('/')).resolve()
                base = skills_dir.resolve()
                if ((folder == base) or (base in folder.parents)) and folder.is_dir():
                    shutil.rmtree(folder, ignore_errors=True)
        conn.execute(f"DELETE FROM prompt_library WHERE id IN ({placeholders})", ids)
        conn.commit()
        return {'deleted': len(rows)}
    finally:
        conn.close()


def open_skill_folder(pid):
    """打开 skill 记录对应的本地文件夹（仅限数据根内，安全约束）。"""
    rec = get_prompt(pid)
    if not rec:
        raise ValueError('记录不存在')
    if rec.get('category') != 'skills' or not rec.get('source_file'):
        raise ValueError('该记录不是 skill，无法打开文件夹')
    folder = (db.DATA_DIR / str(rec['source_file']).lstrip('/')).resolve()
    base = (db.DATA_DIR / SKILLS_DIR_NAME).resolve()
    if not folder.is_dir() or (folder != base and base not in folder.parents):
        raise ValueError('skill 文件夹不存在或超出数据根')
    try:
        os.startfile(str(folder))  # Windows
    except AttributeError:
        try:
            subprocess.run(['xdg-open', str(folder)], check=False)
        except Exception:
            pass
    return True

